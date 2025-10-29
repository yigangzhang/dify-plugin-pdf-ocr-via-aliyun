"""
Test configuration and fixtures for Smart Document Parser Plugin tests.
"""
import sys
import os
from pathlib import Path

# Add project root to Python path for imports
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

import pytest
import io
import base64
from unittest.mock import Mock, MagicMock
from dify_plugin.entities.tool import ToolInvokeMessage

# Import the tool module (handle hyphenated filename)
# This allows tests to use: from tools.smart_doc_parser import SmartDocParserTool
import importlib.util

# Ensure tools package exists
import tools

# Create a virtual module for tools.smart_doc_parser
smart_doc_parser_path = project_root / "tools" / "smart-doc-parser.py"
if smart_doc_parser_path.exists():
    # Load the module with hyphenated filename
    spec = importlib.util.spec_from_file_location(
        "tools.smart_doc_parser", 
        smart_doc_parser_path
    )
    if spec and spec.loader:
        smart_doc_parser_module = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(smart_doc_parser_module)
        # Make it importable as tools.smart_doc_parser
        tools.smart_doc_parser = smart_doc_parser_module
        sys.modules["tools.smart_doc_parser"] = smart_doc_parser_module


@pytest.fixture
def mock_runtime():
    """Mock runtime with credentials."""
    from dify_plugin.entities.tool import ToolRuntime
    
    return ToolRuntime(
        credentials={
            "api_key": "test-api-key",
            "base_url": "https://test.aliyuncs.com/compatible-mode/v1",
            "model": "qwen-vl-ocr",
            "file_host_base": "https://test.example.com"
        },
        user_id="test-user",
        session_id="test-session"
    )


@pytest.fixture
def mock_session():
    """Mock session for tool initialization."""
    from dify_plugin.core.runtime import Session
    
    return Session.empty_session()


@pytest.fixture
def tool_instance(mock_runtime, mock_session):
    """Create a SmartDocParserTool instance for testing."""
    from tools.smart_doc_parser import SmartDocParserTool
    
    tool = SmartDocParserTool(runtime=mock_runtime, session=mock_session)
    return tool


@pytest.fixture
def sample_image_bytes():
    """Sample image bytes for testing."""
    # Create a simple PNG image in bytes (1x1 pixel)
    png_data = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8/5+hHgAHggJ/PchI7wAAAABJRU5ErkJggg=="
    )
    return png_data


@pytest.fixture
def sample_pdf_text_bytes():
    """Sample text-based PDF bytes for testing."""
    # Minimal PDF with text content
    pdf_content = b"""%PDF-1.4
1 0 obj
<<
/Type /Catalog
/Pages 2 0 R
>>
endobj
2 0 obj
<<
/Type /Pages
/Kids [3 0 R]
/Count 1
>>
endobj
3 0 obj
<<
/Type /Page
/Parent 2 0 R
/MediaBox [0 0 612 792]
/Contents 4 0 R
/Resources <<
/Font <<
/F1 5 0 R
>>
>>
>>
endobj
4 0 obj
<<
/Length 44
>>
stream
BT
/F1 12 Tf
100 700 Td
(Hello World!) Tj
ET
endstream
endobj
5 0 obj
<<
/Type /Font
/Subtype /Type1
/BaseFont /Helvetica
>>
endobj
xref
0 6
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000274 00000 n 
0000000368 00000 n 
trailer
<<
/Size 6
/Root 1 0 R
>>
startxref
465
%%EOF"""
    return pdf_content


@pytest.fixture
def sample_pdf_scanned_bytes():
    """Sample scanned PDF bytes (minimal PDF without extractable text)."""
    # Minimal PDF without text content (would be detected as scanned)
    pdf_content = b"""%PDF-1.4
1 0 obj
<<
/Type /Catalog
/Pages 2 0 R
>>
endobj
2 0 obj
<<
/Type /Pages
/Kids [3 0 R]
/Count 1
>>
endobj
3 0 obj
<<
/Type /Page
/Parent 2 0 R
/MediaBox [0 0 612 792]
>>
endobj
xref
0 4
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
trailer
<<
/Size 4
/Root 1 0 R
>>
startxref
180
%%EOF"""
    return pdf_content


@pytest.fixture
def sample_docx_bytes():
    """Sample DOCX bytes for testing."""
    # Create a simple DOCX file in memory
    from io import BytesIO
    try:
        from docx import Document
        doc = Document()
        doc.add_paragraph("This is a test document.")
        doc.add_paragraph("It contains multiple paragraphs.")
        
        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
    except ImportError:
        # If python-docx is not available, return mock ZIP-like data
        return b"PK\x03\x04[Content_Types].xml"


@pytest.fixture
def sample_doc_bytes():
    """Sample DOC bytes for testing (legacy format)."""
    # DOC file magic signature
    return b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1' + b'Microsoft Word Document' + b'\x00' * 100


@pytest.fixture
def mock_openai_client():
    """Mock OpenAI client for testing OCR calls."""
    client = Mock()
    response = Mock()
    choice = Mock()
    message = Mock()
    
    message.content = '{"extracted_text": "Sample OCR result", "confidence": 0.95}'
    choice.message = message
    response.choices = [choice]
    
    client.chat.completions.create.return_value = response
    return client


@pytest.fixture
def sample_tool_parameters():
    """Sample tool parameters for testing."""
    return {
        "prompt": "Extract all text content from this document",
        "file_url": "https://test.example.com/sample.pdf",
        "api_key": "test-override-key",
        "model": "qwen-vl-ocr"
    }


@pytest.fixture
def mock_requests_response():
    """Mock requests response for file downloads."""
    response = Mock()
    response.raise_for_status.return_value = None
    response.headers = {"Content-Type": "application/pdf"}
    return response


@pytest.fixture
def tool_invoke_messages():
    """Factory for creating ToolInvokeMessage objects."""
    def _create_message(type_="text", message="", data=None):
        return ToolInvokeMessage(type=type_, message=message, data=data)
    return _create_message
