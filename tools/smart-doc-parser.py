from collections.abc import Generator
from typing import Any, List, Optional, Tuple
import io
import base64
import mimetypes
import requests
from urllib.parse import urlparse

from dify_plugin import Tool
from dify_plugin.entities.tool import ToolInvokeMessage

from openai import OpenAI
import pypdfium2 as pdfium
from PIL import Image

# Optional imports with fallback handling
try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    from docx import Document  # python-docx
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False

try:
    import textract  # For legacy DOC file processing (optional, requires system dependencies)
    HAS_TEXTRACT = True
except ImportError:
    HAS_TEXTRACT = False

try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False


class SmartDocParserTool(Tool):
    """
    Smart Document Parser Tool for Dify Plugin System
    
    Automatically detects file types and uses appropriate processing:
    - Images (PNG, JPEG, JPG): OCR processing
    - Scanned PDFs: OCR processing  
    - Standard PDFs: Direct text extraction
    - Word documents (DOCX): Direct text extraction using python-docx
    - Legacy Word documents (DOC): Direct text extraction using textract, or OCR fallback
    """
    
    SUPPORTED_IMAGE_TYPES = {'.png', '.jpg', '.jpeg', '.gif', '.webp', '.bmp'}
    SUPPORTED_PDF_TYPES = {'.pdf'}
    SUPPORTED_DOCX_TYPES = {'.docx'}
    SUPPORTED_DOC_TYPES = {'.doc'}
    
    def _invoke(self, tool_parameters: dict[str, Any]) -> Generator[ToolInvokeMessage]:
        prompt: str = str(tool_parameters.get("prompt") or "").strip()
        raw_file_value: Any = tool_parameters.get("file_url")
        file_url: str = self._extract_file_url(raw_file_value)

        # Build absolute file URL from env/provider if needed
        auto_base = self._get_auto_base_url()
        file_url = self._absolutize_url(file_url, auto_base)

        if not prompt:
            yield self.create_text_message("Missing required parameter: prompt")
            return
        if not file_url:
            yield self.create_text_message("Missing required parameter: file_url")
            return
            
        # Basic URL validation
        if not (file_url.startswith("http://") or file_url.startswith("https://")):
            yield self.create_json_message({
                "error": "invalid_file_url",
                "detail": "`file_url` must start with http:// or https://",
                "value": file_url,
            })
            return

        try:
            # Download and analyze the file
            file_bytes, detected_type = self._download_and_detect_file(file_url)
            
            if not file_bytes:
                yield self.create_json_message({
                    "error": "download_failed",
                    "detail": "Could not download or read the file",
                })
                return

            # Process based on detected file type
            result = self._process_file_by_type(file_bytes, detected_type, prompt, tool_parameters)
            
            # Return unified result format
            if isinstance(result, dict) and "error" in result:
                yield self.create_json_message(result)
                return
                
            yield self.create_text_message(self._format_text_output(result))
            yield self.create_json_message(result)
            
        except Exception as e:
            yield self.create_json_message({
                "error": "processing_failed",
                "detail": str(e),
            })

    def _download_and_detect_file(self, file_url: str) -> Tuple[Optional[bytes], str]:
        """Download file and detect its type based on content and URL."""
        try:
            # Download file as binary (preserves all character encodings including Chinese)
            # Note: requests handles URL encoding automatically for Chinese characters in URLs
            response = requests.get(file_url, timeout=30)
            response.raise_for_status()
            # Use response.content to get raw bytes (preserves binary data including Chinese text in files)
            file_bytes = response.content
            
            # Detect file type from URL path and content
            detected_type = self._detect_file_type(file_url, file_bytes)
            
            return file_bytes, detected_type
        except Exception:
            return None, "unknown"

    def _detect_file_type(self, file_url: str, file_bytes: bytes) -> str:
        """Detect file type from URL and content analysis."""
        # Check URL extension first
        parsed = urlparse(file_url)
        path_lower = parsed.path.lower()
        
        # Check by file extension
        for ext in self.SUPPORTED_IMAGE_TYPES:
            if path_lower.endswith(ext):
                return "image"
                
        for ext in self.SUPPORTED_PDF_TYPES:
            if path_lower.endswith(ext):
                return "pdf"
                
        for ext in self.SUPPORTED_DOCX_TYPES:
            if path_lower.endswith(ext):
                return "docx"
                
        for ext in self.SUPPORTED_DOC_TYPES:
            if path_lower.endswith(ext):
                return "doc"
        
        # Check MIME type from content
        if file_bytes:
            # PDF magic bytes
            if file_bytes.startswith(b'%PDF-'):
                return "pdf"
            
            # DOCX magic bytes (ZIP signature + docx content)
            if file_bytes.startswith(b'PK\x03\x04'):
                try:
                    # Could be DOCX (which is a ZIP file)
                    if b'word/' in file_bytes[:4096] or b'[Content_Types].xml' in file_bytes[:4096]:
                        return "docx"
                except:
                    pass
            
            # DOC magic bytes (legacy Microsoft Word)
            if file_bytes.startswith(b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1'):
                return "doc"
            
            # Image magic bytes
            image_signatures = [
                (b'\x89PNG\r\n\x1a\n', "image"),  # PNG
                (b'\xff\xd8\xff', "image"),        # JPEG
                (b'GIF8', "image"),                # GIF
                (b'RIFF', "image"),                # WebP (partial)
                (b'BM', "image"),                  # BMP
            ]
            
            for signature, file_type in image_signatures:
                if file_bytes.startswith(signature):
                    return file_type
        
        return "unknown"

    def _process_file_by_type(self, file_bytes: bytes, file_type: str, prompt: str, tool_parameters: dict) -> dict:
        """Route file processing based on detected type."""
        if file_type == "image":
            return self._process_image_with_ocr(file_bytes, prompt, tool_parameters)
        elif file_type == "pdf":
            return self._process_pdf(file_bytes, prompt, tool_parameters)
        elif file_type == "docx":
            return self._process_docx(file_bytes, prompt)
        elif file_type == "doc":
            return self._process_doc(file_bytes, prompt, tool_parameters)
        else:
            return {
                "error": "unsupported_file_type",
                "detail": f"File type '{file_type}' is not supported. Supported types: images, PDF, DOCX, DOC",
            }

    def _process_image_with_ocr(self, file_bytes: bytes, prompt: str, tool_parameters: dict) -> dict:
        """Process image files using OCR."""
        try:
            # Convert bytes to data URL for OCR API
            data_url = self._bytes_to_data_url(file_bytes, "image/jpeg")
            return self._call_ocr_api([data_url], prompt, tool_parameters)
        except Exception as e:
            return {"error": "image_ocr_failed", "detail": str(e)}

    def _process_pdf(self, file_bytes: bytes, prompt: str, tool_parameters: dict) -> dict:
        """Process PDF files - detect if scanned or standard, then handle accordingly."""
        try:
            is_scanned = self._is_pdf_scanned(file_bytes)
            
            if is_scanned:
                # Scanned PDF - use OCR
                return self._process_scanned_pdf_with_ocr(file_bytes, prompt, tool_parameters)
            else:
                # Standard PDF - extract text directly
                return self._extract_text_from_pdf(file_bytes, prompt)
        except Exception as e:
            return {"error": "pdf_processing_failed", "detail": str(e)}

    def _process_docx(self, file_bytes: bytes, prompt: str) -> dict:
        """Process DOCX files by extracting text directly."""
        if not HAS_PYTHON_DOCX:
            return {
                "error": "docx_not_supported",
                "detail": "python-docx library not installed. Cannot process DOCX files."
            }
        
        try:
            return self._extract_text_from_docx(file_bytes, prompt)
        except Exception as e:
            return {"error": "docx_processing_failed", "detail": str(e)}

    def _process_doc(self, file_bytes: bytes, prompt: str, tool_parameters: dict = None) -> dict:
        """Process legacy DOC files by extracting text directly or using OCR fallback."""
        if tool_parameters is None:
            tool_parameters = {}
        
        # Try direct text extraction with textract if available
        if HAS_TEXTRACT:
            try:
                return self._extract_text_from_doc(file_bytes, prompt)
            except Exception as e:
                # If textract fails, fall back to OCR processing
                # Convert DOC to images and process with OCR
                return self._process_doc_with_ocr_fallback(file_bytes, prompt, tool_parameters, str(e))
        else:
            # No textract available, use OCR processing for DOC files
            return self._process_doc_with_ocr_fallback(file_bytes, prompt, tool_parameters, 
                                                       "textract not available")

    def _is_pdf_scanned(self, file_bytes: bytes) -> bool:
        """Determine if PDF is scanned (image-only) or contains extractable text."""
        # Try PyMuPDF first (more reliable)
        if HAS_PYMUPDF:
            try:
                doc = fitz.open(stream=file_bytes, filetype="pdf")
                # Check first few pages for text content
                pages_to_check = min(3, len(doc))
                total_text_length = 0
                
                for page_num in range(pages_to_check):
                    page = doc[page_num]
                    text = page.get_text().strip()
                    total_text_length += len(text)
                
                doc.close()
                
                # If very little text found, consider it scanned
                # Threshold: less than 50 characters across first few pages
                return total_text_length < 50
            except Exception:
                pass
        
        # Fallback to PyPDF2 if available
        if HAS_PYPDF2:
            try:
                import io
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
                total_text_length = 0
                pages_to_check = min(3, len(pdf_reader.pages))
                
                for i in range(pages_to_check):
                    text = pdf_reader.pages[i].extract_text().strip()
                    total_text_length += len(text)
                
                return total_text_length < 50
            except Exception:
                pass
        
        # If no PDF libraries available, assume scanned (safer for OCR)
        return True

    def _extract_text_from_pdf(self, file_bytes: bytes, prompt: str) -> dict:
        """Extract text directly from standard PDF files."""
        if HAS_PYMUPDF:
            try:
                doc = fitz.open(stream=file_bytes, filetype="pdf")
                pages_content = []
                
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    text = page.get_text()
                    
                    # Apply prompt-based processing to extracted text
                    processed_content = self._process_extracted_text(text.strip(), prompt)
                    pages_content.append({
                        "page": page_num + 1,
                        "content": processed_content
                    })
                
                doc.close()
                return {"pages": pages_content, "extraction_method": "direct_pdf_text"}
            except Exception as e:
                return {"error": "pdf_text_extraction_failed", "detail": str(e)}
        
        elif HAS_PYPDF2:
            try:
                import io
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
                pages_content = []
                
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    processed_content = self._process_extracted_text(text.strip(), prompt)
                    pages_content.append({
                        "page": page_num + 1, 
                        "content": processed_content
                    })
                
                return {"pages": pages_content, "extraction_method": "direct_pdf_text"}
            except Exception as e:
                return {"error": "pdf_text_extraction_failed", "detail": str(e)}
        
        return {
            "error": "pdf_library_missing",
            "detail": "No PDF text extraction library available (PyMuPDF or PyPDF2 required)"
        }

    def _extract_text_from_docx(self, file_bytes: bytes, prompt: str) -> dict:
        """Extract text directly from DOCX files."""
        try:
            doc = Document(io.BytesIO(file_bytes))
            
            # Extract all text from paragraphs
            full_text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text_parts.append(paragraph.text.strip())
            
            # Extract text from tables if any
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text_parts.append(cell.text.strip())
            
            full_text = '\n'.join(full_text_parts)
            
            # Process the extracted text based on the prompt
            processed_content = self._process_extracted_text(full_text, prompt)
            
            return {
                "pages": [{"page": 1, "content": processed_content}],
                "extraction_method": "direct_docx_text"
            }
        except Exception as e:
            return {"error": "docx_text_extraction_failed", "detail": str(e)}

    def _extract_text_from_doc(self, file_bytes: bytes, prompt: str) -> dict:
        """Extract text directly from legacy DOC files using textract."""
        try:
            # Save bytes to temporary file for textract processing
            import tempfile
            import os
            
            with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp_file:
                tmp_file.write(file_bytes)
                tmp_file_path = tmp_file.name
            
            try:
                # Extract text using textract (returns bytes)
                textract_result = textract.process(tmp_file_path)
                
                # Handle both bytes and string responses
                # Support various encodings for Chinese characters
                if isinstance(textract_result, bytes):
                    # Try UTF-8 first, then GBK (common Chinese encoding), then fall back to replace
                    for encoding in ['utf-8', 'gbk', 'gb2312', 'big5']:
                        try:
                            full_text = textract_result.decode(encoding)
                            break
                        except (UnicodeDecodeError, LookupError):
                            continue
                    else:
                        # Fallback: decode with UTF-8 and replace errors
                        full_text = textract_result.decode('utf-8', errors='replace')
                else:
                    full_text = str(textract_result)
                
                if not full_text or not full_text.strip():
                    full_text = ""
                
                # Process the extracted text based on the prompt
                processed_content = self._process_extracted_text(full_text.strip(), prompt)
                
                return {
                    "pages": [{"page": 1, "content": processed_content}],
                    "extraction_method": "direct_doc_text"
                }
            finally:
                # Clean up temporary file
                try:
                    os.unlink(tmp_file_path)
                except:
                    pass
                    
        except Exception as e:
            raise Exception(f"Text extraction failed: {str(e)}")

    def _process_doc_with_ocr_fallback(self, file_bytes: bytes, prompt: str, tool_parameters: dict, reason: str) -> dict:
        """Process DOC files using OCR when text extraction is not available."""
        try:
            # Convert DOC to PDF first using PyMuPDF if available, then use OCR
            if HAS_PYMUPDF:
                # Try to convert DOC to PDF using PyMuPDF
                # Note: PyMuPDF can't directly read DOC files, so we'll use OCR approach
                # For DOC files, we treat them as if they need OCR processing
                # In practice, you might want to use LibreOffice to convert DOC to PDF first
                pass
            
            # For DOC files without textract, we recommend converting to PDF or images first
            # For now, return an informative error suggesting OCR or conversion
            return {
                "error": "doc_processing_requires_conversion",
                "detail": f"Cannot process .doc files directly. {reason}. " +
                         "Please convert .doc to .pdf or .docx format, or use OCR processing. " +
                         "Alternatively, install textract with system dependencies (antiword, etc.) for direct processing.",
                "suggestion": "Convert DOC to DOCX/PDF format or use OCR via image conversion"
            }
        except Exception as e:
            return {"error": "doc_processing_failed", "detail": str(e)}

    def _process_extracted_text(self, text: str, prompt: str) -> dict:
        """Process extracted text according to the prompt to structure it appropriately."""
        # For direct text extraction, we can implement basic structured processing
        # This is a simplified version - in practice, you might want to use LLM for structuring
        if not text.strip():
            return {"raw_text": "", "extracted_fields": {}}
        
        # Basic field extraction based on common patterns
        result = {
            "raw_text": text,
            "extracted_fields": self._extract_basic_fields(text, prompt),
            "word_count": len(text.split()),
            "character_count": len(text)
        }
        
        return result

    def _extract_basic_fields(self, text: str, prompt: str) -> dict:
        """Extract basic structured fields from text based on prompt keywords."""
        import re
        
        fields = {}
        text_lower = text.lower()
        prompt_lower = prompt.lower()
        
        # Common field patterns based on prompt keywords
        field_patterns = {
            "email": r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
            "phone": r'(?:\+?1[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})',
            "date": r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b|\b\d{4}[/-]\d{1,2}[/-]\d{1,2}\b',
            "amount": r'\$\s*\d+(?:,\d{3})*(?:\.\d{2})?|\b\d+(?:,\d{3})*(?:\.\d{2})?\s*(?:USD|EUR|GBP)\b',
        }
        
        # Extract based on what's requested in the prompt
        for field_name, pattern in field_patterns.items():
            if field_name in prompt_lower or f"{field_name}s" in prompt_lower:
                matches = re.findall(pattern, text, re.IGNORECASE)
                if matches:
                    fields[field_name] = matches if len(matches) > 1 else matches[0]
        
        return fields

    def _process_scanned_pdf_with_ocr(self, file_bytes: bytes, prompt: str, tool_parameters: dict) -> dict:
        """Process scanned PDF files using OCR (existing logic from original plugin)."""
        try:
            images = self._convert_pdf_to_data_urls(file_bytes)
            if not images:
                return {
                    "error": "pdf_convert_failed",
                    "detail": "No images could be rendered from PDF"
                }
            
            return self._call_ocr_api(images, prompt, tool_parameters)
        except Exception as e:
            return {"error": "scanned_pdf_ocr_failed", "detail": str(e)}

    def _call_ocr_api(self, images: List[str], prompt: str, tool_parameters: dict) -> dict:
        """Call Aliyun OCR API with the provided images."""
        # Get API credentials (existing logic from original plugin)
        override_key = tool_parameters.get("api_key")
        api_key = str((override_key if override_key is not None else self.runtime.credentials.get("api_key")) or "").strip()
        base_url = str(self.runtime.credentials.get("base_url") or "https://dashscope.aliyuncs.com/compatible-mode/v1").strip()
        
        override_model = tool_parameters.get("model")
        model = str((override_model if override_model is not None else self.runtime.credentials.get("model")) or "qwen-vl-ocr").strip()

        client_kwargs: dict[str, Any] = {"api_key": api_key}
        if base_url:
            client_kwargs["base_url"] = base_url

        client = OpenAI(**client_kwargs)

        pages_result: List[dict[str, Any]] = []
        
        for idx, image_data_url in enumerate(images, start=1):
            page_messages = [
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {"type": "image_url", "image_url": image_data_url},
                    ],
                }
            ]
            
            try:
                resp = client.chat.completions.create(
                    model=model,
                    messages=page_messages,
                    response_format={"type": "json_object"},
                    max_tokens=4096,
                )
                
                page_text = "{}"
                if getattr(resp, "choices", None):
                    try:
                        page_text = resp.choices[0].message.content or "{}"
                    except Exception:
                        page_text = "{}"
                        
                pages_result.append({
                    "page": idx, 
                    "content": self.safe_json_loads(page_text)
                })
            except Exception as e:
                pages_result.append({"page": idx, "error": str(e)})

        return {
            "pages": pages_result,
            "extraction_method": "ocr_api"
        }

    # Helper methods (adapted from original plugin)
    def _convert_pdf_to_data_urls(self, pdf_bytes: bytes) -> List[str]:
        """Convert PDF pages to data URLs for OCR processing."""
        data_urls: List[str] = []
        try:
            with io.BytesIO(pdf_bytes) as bio:
                pdf = pdfium.PdfDocument(bio)
                page_count = len(pdf)
                for index in range(page_count):
                    page = pdf[index]
                    bitmap = page.render(scale=2.0).to_pil()
                    if not isinstance(bitmap, Image.Image):
                        continue
                    buf = io.BytesIO()
                    bitmap.save(buf, format="PNG")
                    data_urls.append(f"data:image/png;base64,{base64.b64encode(buf.getvalue()).decode('utf-8')}")
        except Exception:
            pass
        return data_urls

    def _bytes_to_data_url(self, file_bytes: bytes, content_type: str) -> str:
        """Convert bytes to data URL."""
        encoded = base64.b64encode(file_bytes).decode('utf-8')
        return f"data:{content_type};base64,{encoded}"

    @staticmethod
    def safe_json_loads(s: str) -> Any:
        """Safely parse JSON string, preserving Unicode characters (including Chinese)."""
        try:
            import json
            # json.loads handles Unicode properly by default in Python 3
            return json.loads(s)
        except Exception:
            return {"raw": s}

    def _format_text_output(self, result: dict) -> str:
        """Format result as text for Direct Reply binding."""
        try:
            import json as _json
            return _json.dumps(result, ensure_ascii=False)
        except Exception:
            return str(result)

    # URL handling methods (from original plugin)
    def _extract_file_url(self, value: Any) -> str:
        """Extract file URL from various input formats."""
        # Accept direct string
        if isinstance(value, str):
            text = value.strip()
            if not text:
                return ""
            # Try JSON decoding if looks like JSON
            if (text.startswith("{") and text.endswith("}")) or (text.startswith("[") and text.endswith("]")):
                try:
                    import json
                    decoded = json.loads(text)
                    return self._extract_file_url(decoded)
                except Exception:
                    return text
            return text

        # Accept dict with common url fields
        if isinstance(value, dict):
            for key in ("url", "file_url", "image_url", "src", "href", "value"):
                candidate = value.get(key)
                if isinstance(candidate, str) and candidate.strip():
                    return candidate.strip()
                if isinstance(candidate, dict) or isinstance(candidate, list):
                    extracted = self._extract_file_url(candidate)
                    if extracted:
                        return extracted
            return ""

        # Accept list/tuple and take first resolvable
        if isinstance(value, (list, tuple)):
            for item in value:
                extracted = self._extract_file_url(item)
                if extracted:
                    return extracted
            return ""

        return ""

    def _absolutize_url(self, url: str, base_override: str) -> str:
        """Convert relative URL to absolute using base URL."""
        if not url:
            return url
        # Already absolute
        if url.startswith("http://") or url.startswith("https://"):
            return url
        # Strip quotes if mistakenly included
        if (url.startswith('"') and url.endswith('"')) or (url.startswith("'") and url.endswith("'")):
            url = url[1:-1]
        # Ensure a single leading slash
        if not url.startswith("/"):
            url = f"/{url}"

        # Prefer explicit base if provided
        base = base_override
        if not base:
            # Try to use provider-configured file host base
            base = str(self.runtime.credentials.get("file_host_base") or "").strip()

        if base:
            # Remove trailing slash to avoid double slashes
            while base.endswith("/"):
                base = base[:-1]
            return f"{base}{url}"

        return url

    def _get_auto_base_url(self) -> str:
        """Get automatic base URL from environment."""
        try:
            import os
            # Common envs used by Dify plugin daemon / deployments
            for key in (
                "FILES_URL",
                "INTERNAL_FILES_URL",
            ):
                v = os.getenv(key)
                if v and (v.startswith("http://") or v.startswith("https://")):
                    # Remove trailing slash
                    while v.endswith("/"):
                        v = v[:-1]
                    return v

            # Development heuristic: if connected to local plugin-daemon, assume local web at 3000
            remote = os.getenv("REMOTE_INSTALL_URL", "")
            if remote and ("localhost" in remote or "127.0.0.1" in remote):
                return "http://localhost"
        except Exception:
            pass
        return ""
